from docx import Document
from backend.services.docscan.utils.hash_utils import calculate_sha256
import os

def scan_docx(file_path):

    findings = []

    try:

        doc = Document(file_path)

        findings.append({
            "type": "File Name",
            "value": os.path.basename(file_path)
        })

        findings.append({
            "type": "Paragraphs",
            "value": len(doc.paragraphs)
        })

        findings.append({
            "type": "SHA256",
            "value": calculate_sha256(file_path)
        })

        suspicious = [
            "powershell",
            "cmd.exe",
            "wget",
            "curl",
            "shell"
        ]

        full_text = ""

        for para in doc.paragraphs:
            full_text += para.text + "\n"

        for item in suspicious:

            if item.lower() in full_text.lower():

                findings.append({
                    "type": "Suspicious Content",
                    "value": item
                })

    except Exception as e:

        findings.append({
            "type": "Error",
            "value": str(e)
        })

    return findings